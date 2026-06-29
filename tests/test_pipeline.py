import os
import sys
import tempfile
import unittest
import datetime
from types import SimpleNamespace
from pathlib import Path
from unittest.mock import patch

from PIL import Image
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import processor
import reporter


class PipelineTests(unittest.TestCase):
    def test_cluster_images_uses_transitive_connectivity(self):
        images_meta = [
            {"path": "a.jpg", "lat": 0.0, "lon": 0.0, "time": None},
            {"path": "b.jpg", "lat": 0.0, "lon": 0.0005, "time": None},
            {"path": "c.jpg", "lat": 0.0, "lon": 0.0010, "time": None},
        ]

        clusters = processor.cluster_images(images_meta, radius_meters=90.0)

        self.assertEqual(len(clusters), 1)
        self.assertEqual([img["path"] for img in clusters[0]], ["a.jpg", "b.jpg", "c.jpg"])

    def test_generate_word_report_uses_generic_defaults_when_customer_info_missing(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            site_image = tmp_path / "site.jpg"
            Image.new("RGB", (100, 100), color="white").save(site_image)

            site_folder = tmp_path / "site"
            site_folder.mkdir()
            drawing_path = site_folder / "engineering_layout.png"
            Image.new("RGB", (200, 100), color="black").save(drawing_path)

            report_path = tmp_path / "report.docx"
            site_data = [
                {
                    "address": "123 Example St, Test City, TS",
                    "latitude": 41.0,
                    "longitude": -87.0,
                    "folder_path": str(site_folder),
                    "analysis": {
                        "roof_access": "Unknown",
                        "roof_type": "Unknown",
                        "mounting_structures": [],
                        "hardware": [],
                    },
                    "airspace": "Class G",
                    "airfield_info": "Test Field (1.00 km)",
                    "images": [
                        {
                            "filename": "site.jpg",
                            "path": str(site_image),
                            "selected_for_report": True,
                        }
                    ],
                }
            ]

            basemap = Image.new("RGBA", (1332, 714), color=(240, 244, 248, 255))
            with patch.object(reporter, "query_city_boundary", return_value=None), \
                 patch.object(reporter, "_build_tile_basemap", return_value=(basemap, 13)):
                reporter.generate_word_report(site_data, str(report_path))

            self.assertTrue(report_path.exists())
            doc = Document(str(report_path))
            text = "\n".join(p.text for p in doc.paragraphs)
            cell_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
            combined = f"{text}\n{cell_text}"

            self.assertNotIn("Lansing Police Department", combined)

    def test_draw_styled_map_renders_boundary_and_rings(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            map_path = tmp_path / "dfr_site_map.png"
            site_data = [
                {
                    "address": "Zionsville Police Department, 1075, Parkway Drive, Zionsville, Boone County, Indiana, 46077, United States",
                    "latitude": 39.94709384920635,
                    "longitude": -86.27429027777778,
                    "city": "Zionsville",
                    "state": "Indiana",
                },
                {
                    "address": "Zionsville Town Hall, 1100, West Oak Street, Zionsville, Boone County, Indiana, 46077, United States",
                    "latitude": 39.952111590038314,
                    "longitude": -86.27500775862069,
                    "city": "Zionsville",
                    "state": "Indiana",
                },
            ]
            boundary = {
                "type": "Polygon",
                "coordinates": [[
                    [-86.2825, 39.9440],
                    [-86.2670, 39.9440],
                    [-86.2670, 39.9568],
                    [-86.2825, 39.9568],
                    [-86.2825, 39.9440],
                ]],
            }

            basemap = Image.new("RGBA", (1332, 714), color=(233, 238, 245, 255))
            with patch.object(reporter, "query_city_boundary", return_value=boundary), \
                 patch.object(reporter, "_build_tile_basemap", return_value=(basemap, 13)):
                reporter.draw_styled_map(site_data, str(map_path))

            self.assertTrue(map_path.exists())
            with Image.open(map_path) as img:
                self.assertEqual(img.size, (1400, 900))
                colors = img.convert("RGB").getcolors(maxcolors=200000)
                self.assertIsNotNone(colors)
                self.assertGreater(len(colors), 10)

    def test_draw_styled_map_uses_translucent_ring_fill(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            map_path = tmp_path / "dfr_site_map.png"
            site_data = [
                {
                    "address": "123 Main St, Chicago, IL",
                    "site_name": "Main St",
                    "latitude": 41.862644,
                    "longitude": -87.661244,
                    "city": "Chicago",
                    "state": "IL",
                }
            ]

            basemap = Image.new("RGBA", (1332, 714), color=(100, 150, 200, 255))
            with patch.object(reporter, "query_city_boundary", return_value=None), \
                 patch.object(reporter, "_build_tile_basemap", return_value=(basemap, 13)):
                reporter.draw_styled_map(site_data, str(map_path))

            with Image.open(map_path) as img:
                blended_pixel = img.convert("RGB").getpixel((800, 475))
                self.assertNotEqual(blended_pixel, (100, 150, 200))
                self.assertNotEqual(blended_pixel, (217, 79, 67))
                self.assertGreater(blended_pixel[0], 100)
                self.assertGreater(blended_pixel[2], 100)

    def test_extract_exif_gps_uses_exifread_for_heic_metadata(self):
        class DummyRatio:
            def __init__(self, num, den):
                self.num = num
                self.den = den

        class DummyTag:
            def __init__(self, values, text):
                self.values = values
                self.text = text

            def __str__(self):
                return self.text

        fake_tags = {
            "GPS GPSLatitude": DummyTag([DummyRatio(38, 1), DummyRatio(42, 1), DummyRatio(1321, 50)], "GPS GPSLatitude"),
            "GPS GPSLatitudeRef": DummyTag("N", "N"),
            "GPS GPSLongitude": DummyTag([DummyRatio(90, 1), DummyRatio(24, 1), DummyRatio(2497, 50)], "GPS GPSLongitude"),
            "GPS GPSLongitudeRef": DummyTag("W", "W"),
            "EXIF DateTimeOriginal": DummyTag("2026:05:06 08:03:22", "2026:05:06 08:03:22"),
        }

        with tempfile.NamedTemporaryFile(suffix=".HEIC", delete=False) as tmp:
            tmp.write(b"test")
            tmp_path = tmp.name

        try:
            original_exifread = processor.exifread
            processor.exifread = SimpleNamespace(
                process_file=lambda fh, details=False: fake_tags
            )

            lat, lon, captured, altitude = processor._extract_gps_with_exifread(tmp_path)

            self.assertAlmostEqual(lat, 38.70733888888889)
            self.assertAlmostEqual(lon, -90.41387222222222)
            self.assertEqual(captured, datetime.datetime(2026, 5, 6, 8, 3, 22))
            self.assertIsNone(altitude)
        finally:
            processor.exifread = original_exifread
            os.unlink(tmp_path)

    def test_process_and_organize_images_scopes_to_explicit_image_paths(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            source_dir = tmp_path / "raw"
            output_dir = tmp_path / "processed"
            source_dir.mkdir()

            old_image = source_dir / "old_site.jpg"
            new_images = []
            Image.new("RGB", (100, 100), color="red").save(old_image)
            for idx in range(5):
                image_path = source_dir / f"new_upload_{idx + 1}.jpg"
                Image.new("RGB", (100, 100), color="blue").save(image_path)
                new_images.append(image_path)

            original_extract = processor.extract_exif_gps
            original_reverse_geocode = processor.reverse_geocode
            try:
                processor.extract_exif_gps = lambda path: (
                    41.0,
                    -87.0,
                    datetime.datetime(2026, 6, 15, 12, 0, 0),
                    600.0,
                )
                processor.reverse_geocode = lambda lat, lon: "100 Test St, Test City, TS"

                site_data = processor.process_and_organize_images(
                    str(source_dir),
                    str(output_dir),
                    image_paths=[str(path) for path in new_images],
                )
            finally:
                processor.extract_exif_gps = original_extract
                processor.reverse_geocode = original_reverse_geocode

            processed_filenames = [
                img["filename"]
                for site in site_data
                for img in site.get("images", [])
            ]
            self.assertEqual(sorted(processed_filenames), [f"new_upload_{idx}.jpg" for idx in range(1, 6)])
            self.assertFalse(any(output_dir.rglob("old_site.jpg")))

    def test_create_engineering_drawing_writes_output(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            bg_path = tmp_path / "bg.png"
            out_path = tmp_path / "engineering_layout.png"
            Image.new("RGB", (400, 200), color="gray").save(bg_path)

            markers = [
                {"type": "Electric", "label": "120V", "node_x": 20, "node_y": 30, "label_x": 10, "label_y": 20},
                {"type": "RF", "label": "Antenna", "node_x": 60, "node_y": 50, "label_x": 45, "label_y": 40},
            ]

            reporter.create_engineering_drawing(str(bg_path), str(out_path), markers, "Test note")

            self.assertTrue(out_path.exists())


class TestBubblePlacement(unittest.TestCase):
    """Tests for smart bubble placement in engineering drawings."""

    def test_bubble_near_top_left_avoids_clipping(self):
        """A marker near the top-left corner should place its bubble below/right, not off-screen."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            bg_path = tmp_path / "bg.png"
            out_path = tmp_path / "drawing.png"
            Image.new("RGB", (400, 200), color="gray").save(bg_path)

            markers = [
                {"type": "Electric", "label": "20A 110V AC", "node_x": 2, "node_y": 2, "label_x": 2, "label_y": 2},
            ]
            result = reporter.create_engineering_drawing(str(bg_path), str(out_path), markers, "Test")
            self.assertTrue(out_path.exists())
            with Image.open(out_path) as img:
                self.assertEqual(img.size, (1200, 675))

    def test_bubble_near_right_edge_avoids_clipping(self):
        """A marker near the right edge should not place its bubble off the photo area."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            bg_path = tmp_path / "bg.png"
            out_path = tmp_path / "drawing.png"
            Image.new("RGB", (400, 200), color="gray").save(bg_path)

            markers = [
                {"type": "Data", "label": "CAT6 Drop", "node_x": 95, "node_y": 50, "label_x": 95, "label_y": 50},
            ]
            result = reporter.create_engineering_drawing(str(bg_path), str(out_path), markers, "Test")
            self.assertTrue(out_path.exists())

    def test_multiple_markers_same_area_no_crash(self):
        """Multiple markers clustered together should render without errors."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            bg_path = tmp_path / "bg.png"
            out_path = tmp_path / "drawing.png"
            Image.new("RGB", (400, 200), color="gray").save(bg_path)

            markers = [
                {"type": "Electric", "label": "20A 110V", "node_x": 50, "node_y": 50, "label_x": 35, "label_y": 42},
                {"type": "Data", "label": "CAT6", "node_x": 52, "node_y": 52, "label_x": 37, "label_y": 44},
                {"type": "RF", "label": "5GHz", "node_x": 48, "node_y": 48, "label_x": 33, "label_y": 40},
            ]
            result = reporter.create_engineering_drawing(str(bg_path), str(out_path), markers, "Test")
            self.assertTrue(out_path.exists())
            with Image.open(out_path) as img:
                self.assertEqual(img.size, (1200, 675))

    def test_four_corner_markers(self):
        """Markers at all four corners should all render without clipping."""
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            bg_path = tmp_path / "bg.png"
            out_path = tmp_path / "drawing.png"
            Image.new("RGB", (400, 200), color="gray").save(bg_path)

            markers = [
                {"type": "Electric", "label": "TL corner", "node_x": 3, "node_y": 3, "label_x": 3, "label_y": 3},
                {"type": "Data", "label": "TR corner", "node_x": 97, "node_y": 3, "label_x": 97, "label_y": 3},
                {"type": "RF", "label": "BL corner", "node_x": 3, "node_y": 97, "label_x": 3, "label_y": 97},
                {"type": "Unistrut", "label": "BR corner", "node_x": 97, "node_y": 97, "label_x": 97, "label_y": 97},
            ]
            result = reporter.create_engineering_drawing(str(bg_path), str(out_path), markers, "Test")
            self.assertTrue(out_path.exists())


if __name__ == "__main__":
    unittest.main()
