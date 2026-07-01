import os
import io
import json
import re
import datetime
import logging
import math
import tempfile
import zipfile
import time
from xml.sax.saxutils import escape as xml_escape
import requests
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from geopy.distance import geodesic
from PIL import Image, ImageDraw, ImageFont

logger = logging.getLogger(__name__)


class _AirspaceCache:
    """Cache airspace lookups with TTL and connection pooling."""
    def __init__(self, ttl_seconds=3600):
        self.cache = {}
        self.ttl = ttl_seconds
        self.session = requests.Session()

    def get(self, lat, lon):
        key = (round(lat, 4), round(lon, 4))
        if key in self.cache:
            value, timestamp = self.cache[key]
            if time.time() - timestamp < self.ttl:
                return value
            del self.cache[key]
        return None

    def set(self, lat, lon, value):
        key = (round(lat, 4), round(lon, 4))
        self.cache[key] = (value, time.time())


_airspace_cache = _AirspaceCache()


def _load_font(bold=False, size=28):
    """Load a TrueType font, trying Windows Arial first then Linux DejaVu."""
    candidates = (
        ["arialbd.ttf", "Arial Bold.ttf"] if bold
        else ["arial.ttf", "Arial.ttf"]
    ) + [
        # Linux / Streamlit Cloud paths
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold
        else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
    ]
    for name in candidates:
        try:
            return ImageFont.truetype(name, size)
        except (OSError, IOError):
            continue
    return ImageFont.load_default()


def _normalize_date_string(value):
    """Return YYYY-MM-DD for date-like values, or an empty string."""
    if not value:
        return ""
    if isinstance(value, datetime.datetime):
        return value.date().isoformat()
    if isinstance(value, datetime.date):
        return value.isoformat()

    text = str(value).strip()
    if not text:
        return ""

    for fmt in (
        "%Y-%m-%d",
        "%Y:%m:%d",
        "%Y:%m:%d %H:%M:%S",
        "%Y-%m-%d %H:%M:%S",
        "%B %d, %Y",
        "%m/%d/%Y",
    ):
        try:
            parsed = datetime.datetime.strptime(text, fmt)
            return parsed.date().isoformat()
        except ValueError:
            continue

    try:
        return datetime.date.fromisoformat(text[:10]).isoformat()
    except ValueError:
        return ""


def _format_report_date(value):
    """Return a human-readable date string for the report header."""
    normalized = _normalize_date_string(value)
    if normalized:
        return datetime.datetime.strptime(normalized, "%Y-%m-%d").strftime("%B %d, %Y")
    return datetime.date.today().strftime("%B %d, %Y")


def _extract_survey_date(site_data_list=None, candidate_sites=None, customer_info=None):
    """Resolve the survey date from EXIF-derived site data when available."""
    if customer_info:
        normalized = _normalize_date_string(customer_info.get("survey_date"))
        if normalized:
            return normalized

    if candidate_sites:
        normalized = _normalize_date_string(candidate_sites[0].identity.survey_date)
        if normalized:
            return normalized

    date_counts = {}
    for site in site_data_list or []:
        site_date = _normalize_date_string(site.get("survey_date"))
        if site_date:
            date_counts[site_date] = date_counts.get(site_date, 0) + 1
        for img in site.get("images", []):
            capture_time = img.get("time")
            normalized = _normalize_date_string(capture_time)
            if normalized:
                date_counts[normalized] = date_counts.get(normalized, 0) + 1

    if date_counts:
        return sorted(date_counts.items(), key=lambda item: (-item[1], item[0]))[0][0]
    return ""


def _survey_site_records(site_data_list=None, candidate_sites=None):
    """Normalize either legacy site dicts or CandidateSite objects into map records."""
    records = []
    if candidate_sites:
        for idx, site in enumerate(candidate_sites, start=1):
            records.append({
                "latitude": site.identity.site_latitude,
                "longitude": site.identity.site_longitude,
                "address": site.identity.site_address or site.identity.site_name or f"Site {idx}",
                "site_name": site.identity.site_name or f"Site {idx}",
                "city": site.identity.jurisdiction,
                "state": site.identity.state,
            })
        return records

    for idx, site in enumerate(site_data_list or [], start=1):
        address = site.get("address") or f"Site {idx}"
        records.append({
            "latitude": site.get("latitude"),
            "longitude": site.get("longitude"),
            "address": address,
            "site_name": site.get("site_name") or address or f"Site {idx}",
            "city": site.get("city") or getattr(address, "city", None),
            "state": site.get("state") or getattr(address, "state", None),
        })
    return records


def _build_black_logo_copy():
    """Create a black version of the BRINC logo for report footers."""
    logo_path = os.path.join(os.path.dirname(__file__), "images", "BRINC_Logo_White.png")
    if not os.path.exists(logo_path):
        return None

    black_logo_path = os.path.join(tempfile.gettempdir(), "brinc_logo_black.png")
    try:
        source_mtime = os.path.getmtime(logo_path)
        if os.path.exists(black_logo_path) and os.path.getmtime(black_logo_path) >= source_mtime:
            return black_logo_path

        with Image.open(logo_path) as logo_src:
            logo = logo_src.convert("RGBA")
            pixels = logo.load()
            for y in range(logo.height):
                for x in range(logo.width):
                    r, g, b, a = pixels[x, y]
                    if a > 0:
                        pixels[x, y] = (0, 0, 0, a)
            logo.save(black_logo_path)
        return black_logo_path
    except Exception:
        return None


def _add_report_footer(doc):
    """Add the BRINC logo to the first section footer."""
    try:
        footer = doc.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_path = _build_black_logo_copy()
        if logo_path:
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.25))
    except Exception:
        pass


def _add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _site_label(site, idx):
    if isinstance(site, dict):
        return (
            site.get("site_name")
            or site.get("folder_name")
            or site.get("site_id")
            or f"Site {idx}"
        )

    identity = getattr(site, "identity", None)
    if identity is not None:
        return (
            getattr(identity, "site_name", None)
            or getattr(identity, "site_id", None)
            or f"Site {idx}"
        )

    return f"Site {idx}"


def _site_address(site):
    if isinstance(site, dict):
        return site.get("address") or site.get("site_address") or ""

    identity = getattr(site, "identity", None)
    if identity is not None:
        return getattr(identity, "site_address", "") or ""

    return ""


def _site_photos(site):
    if isinstance(site, dict):
        return site.get("images", []) or []

    return getattr(site, "photos", []) or []


def _photo_lookup_keys(photo):
    keys = []
    if isinstance(photo, dict):
        for field in ("filename", "path", "dest_path", "photo_id", "file_path"):
            value = photo.get(field)
            if value:
                keys.append(os.path.basename(str(value)))
                keys.append(str(value))
    else:
        for field in ("photo_id", "file_path"):
            value = getattr(photo, field, None)
            if value:
                keys.append(os.path.basename(str(value)))
                keys.append(str(value))
    return [key for key in keys if key]


def append_drive_photo_links(report_path, photo_links_by_name, site_data_list=None):
    """Append a Drive links section to an already-rendered DOCX report."""
    if not report_path or not os.path.exists(report_path):
        raise FileNotFoundError(f"Report not found: {report_path}")

    photo_links_by_name = photo_links_by_name or {}
    if not photo_links_by_name and not site_data_list:
        return report_path

    doc = Document(report_path)
    if any((p.text or "").strip() == "Drive Photo Links" for p in doc.paragraphs):
        return report_path

    doc.add_page_break()
    doc.add_heading("Drive Photo Links", level=1)
    doc.add_paragraph("Open the uploaded survey photos in Google Drive using the links below.")

    sites = site_data_list or []
    if sites:
        for idx, site in enumerate(sites, start=1):
            site_label = _site_label(site, idx)
            site_address = _site_address(site)
            photos = [
                p for p in _site_photos(site)
                if (p.get("selected_for_report", True) if isinstance(p, dict) else getattr(p, "selected_for_report", True))
            ]
            site_header = doc.add_paragraph()
            site_header.add_run(f"{site_label}").bold = True
            if site_address:
                site_header.add_run(f" - {site_address}")

            for photo in photos:
                link = None
                if isinstance(photo, dict) and photo.get("drive_url"):
                    link = photo["drive_url"]
                for key in _photo_lookup_keys(photo):
                    if key in photo_links_by_name:
                        link = photo_links_by_name[key]
                        break
                if not link:
                    continue

                if isinstance(photo, dict):
                    photo_name = (
                        photo.get("filename")
                        or os.path.basename(photo.get("path", ""))
                        or os.path.basename(photo.get("dest_path", ""))
                    )
                else:
                    photo_name = getattr(photo, "photo_id", "") or os.path.basename(getattr(photo, "file_path", ""))

                p = doc.add_paragraph()
                p.add_run(f"{photo_name}: ")
                _add_hyperlink(p, "Open in Google Drive", link)
    else:
        for photo_name, link in sorted(photo_links_by_name.items()):
            p = doc.add_paragraph()
            p.add_run(f"{photo_name}: ")
            _add_hyperlink(p, "Open in Google Drive", link)

    doc.save(report_path)
    return report_path


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    tc_pr.append(shd)


def _set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_column_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def _style_table_cell(cell, *, bold=False, font_size=10, color=None, align=None):
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(font_size)
            run.font.bold = bold
            if color is not None:
                run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)


def _style_polished_table(table, *, widths, header_fill="1F4E79", label_fill="D9E8FB", body_fill="FFFFFF"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_column_widths(table, widths)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            _set_cell_shading(cell, header_fill if row_idx == 0 else (label_fill if col_idx == 0 else body_fill))
            if row_idx == 0:
                _style_table_cell(cell, bold=True, font_size=10, color=RGBColor(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
            elif col_idx == 0:
                _style_table_cell(cell, bold=True, font_size=10, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                _style_table_cell(cell, bold=False, font_size=10, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.LEFT)


def export_sites_kmz(site_data_list, output_filepath, candidate_sites=None, ring_radius_miles=2.0):
    """Export site points and 2-mile rings to a KMZ file."""
    records = _survey_site_records(site_data_list, candidate_sites)
    doc_name = os.path.splitext(os.path.basename(output_filepath))[0]

    placemarks = []
    for idx, site in enumerate(records, start=1):
        lat = site.get("latitude")
        lon = site.get("longitude")
        if lat is None or lon is None:
            continue
        site_name = xml_escape(str(site.get("site_name") or f"Site {idx}"))
        address = xml_escape(str(site.get("address") or ""))

        ring_coords = []
        for bearing in range(0, 361, 10):
            ring_point = geodesic(miles=ring_radius_miles).destination((lat, lon), bearing)
            ring_coords.append(f"{ring_point.longitude:.8f},{ring_point.latitude:.8f},0")

        placemarks.append(f"""
        <Placemark>
          <name>{site_name}</name>
          <description>{address}</description>
          <Style>
            <IconStyle><color>ff0000ff</color><scale>1.1</scale></IconStyle>
          </Style>
          <Point><coordinates>{lon:.8f},{lat:.8f},0</coordinates></Point>
        </Placemark>
        <Placemark>
          <name>{site_name} - {ring_radius_miles:.0f} Mile Ring</name>
          <Style>
            <LineStyle><color>ff0000ff</color><width>2</width></LineStyle>
            <PolyStyle><color>220000ff</color></PolyStyle>
          </Style>
          <Polygon>
            <outerBoundaryIs>
              <LinearRing>
                <coordinates>
                  {' '.join(ring_coords)}
                </coordinates>
              </LinearRing>
            </outerBoundaryIs>
          </Polygon>
        </Placemark>
        """)

    kml = f"""<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <name>{xml_escape(doc_name)}</name>
    <Style id="sitePoint">
      <IconStyle>
        <color>ff0000ff</color>
        <scale>1.1</scale>
      </IconStyle>
    </Style>
    <Style id="ringStyle">
      <LineStyle>
        <color>ff0000ff</color>
        <width>2</width>
      </LineStyle>
      <PolyStyle>
        <color>220000ff</color>
      </PolyStyle>
    </Style>
    {''.join(placemarks)}
  </Document>
</kml>
"""

    kml_name = os.path.splitext(os.path.basename(output_filepath))[0] + ".kml"
    with zipfile.ZipFile(output_filepath, "w", zipfile.ZIP_DEFLATED) as kmz:
        kmz.writestr(kml_name, kml)
    return output_filepath


def _format_short_address(full_address):
    """Shorten a Nominatim address to just street, city, state zip."""
    if not full_address:
        return str(full_address)
    addr = str(full_address).strip()
    if addr.startswith("Site Coordinate"):
        return addr

    parts = [p.strip() for p in addr.split(',') if p.strip()]

    # Remove country names and county/parish parts
    excluded = {'united states', 'usa', 'us'}
    filtered = [
        p for p in parts
        if p.lower() not in excluded
        and 'county' not in p.lower()
        and 'parish' not in p.lower()
    ]

    # Drop leading POI/building name (e.g. "Polk Avenue Police Department")
    if len(filtered) >= 2 and not filtered[0][0].isdigit():
        first_lower = filtered[0].lower()
        poi_suffixes = (
            'police department', 'fire department', 'fire station',
            'city hall', 'courthouse', 'school', 'university',
            'hospital', 'library', 'church', 'center', 'centre',
            'office', 'building', 'station', 'academy', 'institute',
        )
        is_poi = any(first_lower.endswith(suffix) for suffix in poi_suffixes)
        # Also treat as POI if followed by a street address (starts with digit)
        if not is_poi and filtered[1].strip()[:1].isdigit():
            is_poi = True
        if is_poi:
            filtered = filtered[1:]

    # Merge house number + street name  ("1075", "Parkway Drive" -> "1075 Parkway Drive")
    combined = []
    i = 0
    while i < len(filtered):
        if (filtered[i].isdigit()
                and i + 1 < len(filtered)
                and not filtered[i + 1][0].isdigit()):
            combined.append(f"{filtered[i]} {filtered[i + 1]}")
            i += 2
        else:
            combined.append(filtered[i])
            i += 1

    # Merge state + zip  ("Indiana", "46077" -> "Indiana 46077")
    final = []
    i = 0
    while i < len(combined):
        if (i + 1 < len(combined)
                and re.match(r'^\d{5}(-\d{4})?$', combined[i + 1])):
            final.append(f"{combined[i]} {combined[i + 1]}")
            i += 2
        else:
            final.append(combined[i])
            i += 1

    return ', '.join(final)

def query_nearest_airfield(lat, lon):
    """
    Search for the nearest aerodrome (airfield/airport) within 25km
    using the OpenStreetMap Overpass API.
    """
    url = "https://overpass-api.de/api/interpreter"
    query = f"""
    [out:json][timeout:5];
    (
      node["aeroway"="aerodrome"](around:25000,{lat},{lon});
      way["aeroway"="aerodrome"](around:25000,{lat},{lon});
      relation["aeroway"="aerodrome"](around:25000,{lat},{lon});
    );
    out center;
    """
    try:
        response = requests.post(url, data={"data": query}, timeout=1.5)
        if response.status_code == 200:
            data = response.json()
            elements = data.get("elements", [])
            
            nearest = None
            min_dist = float('inf')
            
            for elem in elements:
                elem_lat = elem.get('lat') or elem.get('center', {}).get('lat')
                elem_lon = elem.get('lon') or elem.get('center', {}).get('lon')
                
                if elem_lat and elem_lon:
                    dist = geodesic((lat, lon), (elem_lat, elem_lon)).miles
                    if dist < min_dist:
                        min_dist = dist
                        tags = elem.get("tags", {})
                        name = tags.get("name") or tags.get("icao") or tags.get("iata") or "Unnamed Airfield"
                        nearest = (name, min_dist)
            return nearest
    except Exception as e:
        logger.warning("Overpass nearest-airfield query failed: %s", e)

    return None

def query_airspace_class(lat, lon):
    """
    Lookup Airspace Class (B, C, D, E, G) for coordinate.
    Uses caching, connection pooling, and retry logic for robustness.
    """
    cached = _airspace_cache.get(lat, lon)
    if cached is not None:
        return cached

    faa_url = "https://services6.arcgis.com/ssFJjBXIUyZDrSYZ/arcgis/rest/services/Class_Airspace/FeatureServer/0/query"
    params = {
        "geometry": json.dumps({
            "x": lon,
            "y": lat,
            "spatialReference": {"wkid": 4326},
        }),
        "geometryType": "esriGeometryPoint",
        "inSR": "4326",
        "spatialRel": "esriSpatialRelIntersects",
        "outFields": "CLASS",
        "f": "json",
        "returnGeometry": "false"
    }

    for attempt in range(3):
        try:
            response = _airspace_cache.session.get(faa_url, params=params, timeout=10)
            if response.status_code == 200:
                data = response.json()
                features = data.get("features", [])
                classes = []
                for feature in features:
                    airspace_class = feature.get("attributes", {}).get("CLASS")
                    if airspace_class and airspace_class not in classes:
                        classes.append(airspace_class)

                if classes:
                    class_order = {"A": 0, "B": 1, "C": 2, "D": 3, "E": 4, "G": 5}
                    classes.sort(key=lambda c: class_order.get(c, 99))
                    if len(classes) == 1:
                        result = f"Controlled (Class {classes[0]})"
                    else:
                        result = f"Controlled (Classes {', '.join(classes)})"
                else:
                    result = "Uncontrolled (Class G)"

                _airspace_cache.set(lat, lon, result)
                return result

            elif response.status_code >= 500 and attempt < 2:
                wait = 2 ** attempt
                logger.debug("FAA server error, retry in %ds", wait)
                time.sleep(wait)
                continue

        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            if attempt < 2:
                wait = 2 ** attempt
                logger.debug("FAA lookup timeout/connection error, retry in %ds", wait)
                time.sleep(wait)
                continue
            logger.warning("FAA airspace lookup failed after retries: %s", e)
        except Exception as e:
            logger.warning("FAA airspace lookup failed: %s", e)

    return None


def query_city_boundary(city_name, state_name=None):
    """Fetch the GeoJSON boundary polygon for a city from OpenStreetMap.

    Uses the Overpass API to find the administrative boundary matching the
    city name. Tries admin_level 8 first (city/town in the US), then falls
    back to levels 7 and 6. If Overpass returns nothing, tries Nominatim's
    structured search as a last resort.

    Args:
        city_name: e.g. "Zionsville"
        state_name: e.g. "Indiana" (optional, helps disambiguate)

    Returns:
        dict with GeoJSON geometry (type, coordinates), or None on failure.
    """
    if not city_name:
        return None

    # Try Overpass with multiple admin levels
    # Prefer admin_level 8 (town/city), then 7 (city), then 6 (county/municipality)
    for admin_level in ("8", "7", "6"):
        result = _query_overpass_boundary(city_name, state_name, admin_level)
        if result and _is_reasonable_city_boundary(result):
            return result

    # Fallback: Nominatim search for boundary polygon
    return _query_nominatim_boundary(city_name, state_name)


def _is_reasonable_city_boundary(boundary_geom):
    """Validate that a boundary represents a city, not a larger region.

    Checks that the boundary fits within reasonable dimensions for a small-to-medium city.
    Small towns: 1-5km. Medium cities: 5-20km. Rejects anything larger.
    """
    try:
        from shapely.geometry import shape
        geom = shape(boundary_geom)
        bounds = geom.bounds  # (minlon, minlat, maxlon, maxlat)

        # Calculate extent in km (rough conversion: 1° ≈ 111km)
        lat_span_km = (bounds[3] - bounds[1]) * 111
        lon_span_km = (bounds[2] - bounds[0]) * 111 * 0.7  # Adjust for latitude

        # Strict: for US small towns/cities, max extent should be 15km max
        # (Zionsville is ~2km across, Indianapolis is ~30km)
        max_extent = max(lat_span_km, lon_span_km)

        return max_extent < 15
    except Exception:
        return True  # If validation fails, accept the boundary


def _query_overpass_boundary(city_name, state_name, admin_level):
    """Query the Overpass API for a city boundary at a given admin level."""
    url = "https://overpass-api.de/api/interpreter"
    area_filter = ""
    if state_name:
        area_filter = f'area["name"="{state_name}"]["admin_level"="4"]->.state;'

    in_area = "(area.state)" if state_name else ""
    query = f"""
    [out:json][timeout:10];
    {area_filter}
    relation["name"="{city_name}"]["admin_level"="{admin_level}"]{in_area};
    out geom;
    """
    try:
        headers = {"Accept": "application/json", "User-Agent": "DFR-SiteSurvey/1.0"}
        response = requests.post(url, data={"data": query}, headers=headers, timeout=20)
        if response.status_code != 200:
            return None
        data = response.json()
        elements = data.get("elements", [])
        if not elements:
            return None

        relation = elements[0]
        outer_rings = []
        for member in relation.get("members", []):
            if member.get("type") == "way" and member.get("role") in ("outer", ""):
                coords = [(pt["lon"], pt["lat"]) for pt in member.get("geometry", [])]
                if coords:
                    outer_rings.append(coords)

        if not outer_rings:
            return None

        merged = _merge_way_segments(outer_rings)

        if len(merged) == 1:
            return {"type": "Polygon", "coordinates": [merged[0]]}
        else:
            return {"type": "MultiPolygon", "coordinates": [[ring] for ring in merged]}

    except Exception as e:
        logger.warning("Overpass boundary query failed (admin_level=%s): %s", admin_level, e)
        return None


def _query_nominatim_boundary(city_name, state_name):
    """Fallback: use Nominatim search to get a city boundary polygon."""
    try:
        params = {
            "city": city_name,
            "format": "geojson",
            "polygon_geojson": 1,
            "limit": 1,
        }
        if state_name:
            params["state"] = state_name
        headers = {"User-Agent": "DFR-SiteSurvey/1.0"}
        resp = requests.get(
            "https://nominatim.openstreetmap.org/search",
            params=params,
            headers=headers,
            timeout=15,
        )
        if resp.status_code != 200:
            return None
        data = resp.json()
        features = data.get("features", [])
        if not features:
            return None
        geom = features[0].get("geometry")
        if geom and geom.get("type") in ("Polygon", "MultiPolygon"):
            return geom
    except Exception as e:
        logger.warning("Nominatim boundary fallback failed: %s", e)
    return None


def _merge_way_segments(segments):
    """Merge ordered way segments into closed rings by connecting endpoints."""
    if not segments:
        return []
    # Work with copies
    remaining = [list(s) for s in segments]
    rings = []

    while remaining:
        current = remaining.pop(0)
        changed = True
        while changed:
            changed = False
            for i, seg in enumerate(remaining):
                # Check if seg connects to the end of current
                if _coords_close(current[-1], seg[0]):
                    current.extend(seg[1:])
                    remaining.pop(i)
                    changed = True
                    break
                elif _coords_close(current[-1], seg[-1]):
                    current.extend(reversed(seg[:-1]))
                    remaining.pop(i)
                    changed = True
                    break
                elif _coords_close(current[0], seg[-1]):
                    current = seg[:-1] + current
                    remaining.pop(i)
                    changed = True
                    break
                elif _coords_close(current[0], seg[0]):
                    current = list(reversed(seg[1:])) + current
                    remaining.pop(i)
                    changed = True
                    break
        # Close the ring if needed
        if not _coords_close(current[0], current[-1]):
            current.append(current[0])
        rings.append(current)

    return rings


def _coords_close(a, b, tol=1e-6):
    """Check if two (lon, lat) coordinate pairs are approximately equal."""
    return abs(a[0] - b[0]) < tol and abs(a[1] - b[1]) < tol


def _draw_dashed_line(draw, points, color, width=2, dash_length=6, gap_length=4):
    """Draw a dashed line using a series of short segments, maintaining pattern across edges."""
    if len(points) < 2:
        return

    dash_pos = 0
    drawing = True
    cycle_length = dash_length + gap_length

    for i in range(len(points) - 1):
        x1, y1 = points[i]
        x2, y2 = points[i + 1]

        dx = x2 - x1
        dy = y2 - y1
        dist = math.sqrt(dx * dx + dy * dy)

        if dist == 0:
            continue

        edge_pos = 0
        while edge_pos < dist:
            phase = dash_pos % cycle_length
            is_drawing = phase < dash_length
            remaining_in_phase = (dash_length if is_drawing else gap_length) - (dash_pos % (dash_length if is_drawing else gap_length))
            segment_length = min(remaining_in_phase, dist - edge_pos)

            t_start = edge_pos / dist
            t_end = (edge_pos + segment_length) / dist

            segment_x1 = x1 + dx * t_start
            segment_y1 = y1 + dy * t_start
            segment_x2 = x1 + dx * t_end
            segment_y2 = y1 + dy * t_end

            if is_drawing:
                draw.line([(segment_x1, segment_y1), (segment_x2, segment_y2)], fill=color, width=width)

            edge_pos += segment_length
            dash_pos += segment_length


def _infer_city_state(site_data_list):
    """Infer a shared city/state label for the exported map from site records."""
    if not site_data_list:
        return None, None

    for site in site_data_list:
        city = (site.get("city") or "").strip()
        state = (site.get("state") or "").strip()
        if city:
            return city, state or None

    for site in site_data_list:
        address = str(site.get("address") or "")
        parts = [part.strip() for part in address.split(",") if part.strip()]
        if len(parts) >= 5:
            state = parts[-3]
            city_idx = -5 if "county" in parts[-4].lower() else -4
            city = parts[city_idx]
            if city and state and state.lower() not in {"united states", "usa"}:
                return city, state
        if len(parts) >= 3:
            city = parts[1] if len(parts) > 3 else parts[0]
            state = parts[-2].split()[0]
            if city and state and state.lower() not in {"united states", "usa", "county"}:
                return city, state

    return None, None


def _infer_boundary_label(site_data_list):
    """Infer the boundary name and mode for a site set."""
    if not site_data_list:
        return None, None

    for site in site_data_list:
        mode = (site.get("jurisdiction_mode") or "").strip().lower()
        label = (site.get("jurisdiction_name") or "").strip()
        state = (site.get("state") or "").strip()
        if mode in {"county", "parish"} and label:
            return label, state or None

    city, state = _infer_city_state(site_data_list)
    return city, state


def _geojson_outer_rings(geometry):
    """Return outer rings from a GeoJSON polygon or multipolygon."""
    if not geometry:
        return []
    if geometry.get("type") == "Polygon":
        coords = geometry.get("coordinates") or []
        return [coords[0]] if coords else []
    if geometry.get("type") == "MultiPolygon":
        rings = []
        for polygon in geometry.get("coordinates") or []:
            if polygon:
                rings.append(polygon[0])
        return rings
    return []


def _project_lon_lat(lon, lat, origin_lon, origin_lat):
    """Project lon/lat to local meters using an equirectangular approximation."""
    meters_per_deg_lat = 111320.0
    meters_per_deg_lon = meters_per_deg_lat * math.cos(math.radians(origin_lat))
    x = (lon - origin_lon) * meters_per_deg_lon
    y = (lat - origin_lat) * meters_per_deg_lat
    return x, y


def _map_site_labels(site_data_list):
    """Build distinct, readable labels for all site markers."""
    names = [(site.get("site_name") or "").strip() for site in site_data_list]
    counts = {}
    for name in names:
        if name:
            key = name.lower()
            counts[key] = counts.get(key, 0) + 1

    labels = []
    for idx, site in enumerate(site_data_list, start=1):
        name = (site.get("site_name") or "").strip()
        city = (site.get("city") or "").strip().lower()
        address = str(site.get("address") or "").strip()
        address_label = address.split(",")[0].strip() if address else ""
        use_address = (
            not name
            or counts.get(name.lower(), 0) > 1
            or (city and name.lower() == city)
        )
        label = address_label if use_address and address_label else name or address_label or f"Site {idx}"
        labels.append(label)
    return labels


def _latlon_to_world_pixels(lat, lon, zoom):
    """Convert WGS84 coordinates to Web Mercator world pixels for a zoom level."""
    lat = max(min(lat, 85.05112878), -85.05112878)
    scale = 256 * (2 ** zoom)
    x = (lon + 180.0) / 360.0 * scale
    sin_lat = math.sin(math.radians(lat))
    y = (
        0.5
        - math.log((1 + sin_lat) / (1 - sin_lat)) / (4 * math.pi)
    ) * scale
    return x, y


def _choose_basemap_zoom(min_lat, max_lat, min_lon, max_lon, width_px, height_px, max_zoom=16):
    """Select the most detailed zoom that keeps tile fetch counts reasonable."""
    if width_px <= 0 or height_px <= 0:
        return 12

    for zoom in range(max_zoom, 0, -1):
        left_x, bottom_y = _latlon_to_world_pixels(min_lat, min_lon, zoom)
        right_x, top_y = _latlon_to_world_pixels(max_lat, max_lon, zoom)
        left = min(left_x, right_x)
        right = max(left_x, right_x)
        top = min(top_y, bottom_y)
        bottom = max(top_y, bottom_y)

        tile_x0 = int(math.floor(left / 256))
        tile_x1 = int(math.floor((right - 1) / 256))
        tile_y0 = int(math.floor(top / 256))
        tile_y1 = int(math.floor((bottom - 1) / 256))
        tile_count = (tile_x1 - tile_x0 + 1) * (tile_y1 - tile_y0 + 1)

        if tile_count <= 36:
            return zoom

    return 1


def _fetch_map_tile(x, y, zoom):
    """Fetch a light street-map tile and return it as an RGBA PIL image."""
    max_tile = 2 ** zoom
    if y < 0 or y >= max_tile:
        return None

    wrapped_x = x % max_tile
    tile_url = f"https://a.basemaps.cartocdn.com/light_all/{zoom}/{wrapped_x}/{y}.png"
    response = requests.get(
        tile_url,
        timeout=8,
        headers={"User-Agent": "DFR-SiteSurvey/1.0"},
    )
    response.raise_for_status()
    return Image.open(io.BytesIO(response.content)).convert("RGBA")


def _build_tile_basemap(min_lat, max_lat, min_lon, max_lon, width_px, height_px):
    """Render a muted street basemap for the requested geographic extent."""
    zoom = _choose_basemap_zoom(min_lat, max_lat, min_lon, max_lon, width_px, height_px)
    min_world_x, max_world_y = _latlon_to_world_pixels(min_lat, min_lon, zoom)
    max_world_x, min_world_y = _latlon_to_world_pixels(max_lat, max_lon, zoom)

    left = min(min_world_x, max_world_x)
    right = max(min_world_x, max_world_x)
    top = min(min_world_y, max_world_y)
    bottom = max(min_world_y, max_world_y)

    tile_size = 256
    tile_x0 = int(math.floor(left / tile_size))
    tile_x1 = int(math.floor((right - 1) / tile_size))
    tile_y0 = int(math.floor(top / tile_size))
    tile_y1 = int(math.floor((bottom - 1) / tile_size))

    stitched = Image.new(
        "RGBA",
        ((tile_x1 - tile_x0 + 1) * tile_size, (tile_y1 - tile_y0 + 1) * tile_size),
        (244, 247, 251, 255),
    )

    for tile_x in range(tile_x0, tile_x1 + 1):
        for tile_y in range(tile_y0, tile_y1 + 1):
            try:
                tile = _fetch_map_tile(tile_x, tile_y, zoom)
            except Exception:
                tile = None
            if tile is not None:
                stitched.paste(tile, ((tile_x - tile_x0) * tile_size, (tile_y - tile_y0) * tile_size))

    crop_left = int(round(left - (tile_x0 * tile_size)))
    crop_top = int(round(top - (tile_y0 * tile_size)))
    crop_right = int(round(right - (tile_x0 * tile_size)))
    crop_bottom = int(round(bottom - (tile_y0 * tile_size)))
    cropped = stitched.crop((crop_left, crop_top, crop_right, crop_bottom))
    return cropped.resize((width_px, height_px), Image.Resampling.LANCZOS), zoom


def _map_extent_with_rings(valid_sites, boundary_rings, ring_radius_m):
    """Compute a padded lat/lon extent that includes site rings and any boundary geometry."""
    lats = []
    lons = []

    for site in valid_sites:
        lat = site["latitude"]
        lon = site["longitude"]
        lats.append(lat)
        lons.append(lon)
        for bearing in (0, 90, 180, 270):
            edge = geodesic(meters=ring_radius_m).destination((lat, lon), bearing)
            lats.append(edge.latitude)
            lons.append(edge.longitude)

    min_lat = min(lats)
    max_lat = max(lats)
    min_lon = min(lons)
    max_lon = max(lons)

    lat_pad = max((max_lat - min_lat) * 0.12, 0.0045)
    lon_pad = max((max_lon - min_lon) * 0.12, 0.0045)
    return (
        max(-85.05112878, min_lat - lat_pad),
        min(85.05112878, max_lat + lat_pad),
        max(-180.0, min_lon - lon_pad),
        min(180.0, max_lon + lon_pad),
    )


def draw_styled_map(site_data_list, output_map_path):
    """
    Create a geographic static map with site markers, 2-mile rings, and city boundary.
    """
    width, height = 1400, 900
    img = Image.new("RGBA", (width, height), color="#eef4fb")
    draw = ImageDraw.Draw(img)

    header_h = 118
    footer_h = 68
    panel_margin = 34
    map_left = panel_margin
    map_top = header_h
    map_right = width - panel_margin
    map_bottom = height - footer_h
    draw.rectangle([0, 0, width, header_h], fill="#0f2743")
    draw.rectangle([0, header_h, width, height], fill="#eef4fb")
    draw.rounded_rectangle(
        [map_left, map_top, map_right, map_bottom],
        radius=18,
        fill="#f8fbff",
        outline="#c7d7ea",
        width=2,
    )

    title_font = _load_font(bold=True, size=34)
    subtitle_font = _load_font(size=19)
    label_font = _load_font(bold=True, size=20)
    small_font = _load_font(size=16)
    marker_font = _load_font(bold=True, size=14)

    if not site_data_list:
        draw.text((40, 34), "DFR SITE MAP", fill="#ffffff", font=title_font)
        draw.text((40, 76), "No survey sites available for rendering.", fill="#d8e4f2", font=subtitle_font)
        img.save(output_map_path)
        return output_map_path

    valid_sites = [
        site for site in site_data_list
        if site.get("latitude") is not None and site.get("longitude") is not None
    ]
    if not valid_sites:
        draw.text((40, 34), "DFR SITE MAP", fill="#ffffff", font=title_font)
        draw.text((40, 76), "Survey sites are missing coordinates.", fill="#d8e4f2", font=subtitle_font)
        img.save(output_map_path)
        return output_map_path

    boundary_name, state = _infer_boundary_label(valid_sites)
    site_labels = _map_site_labels(valid_sites)
    boundary = query_city_boundary(boundary_name, state) if boundary_name else None

    ring_radius_m = 3218.69
    boundary_rings = _geojson_outer_rings(boundary)
    map_w = map_right - map_left
    map_h = map_bottom - map_top
    min_lat, max_lat, min_lon, max_lon = _map_extent_with_rings(valid_sites, boundary_rings, ring_radius_m)
    basemap, zoom = _build_tile_basemap(min_lat, max_lat, min_lon, max_lon, map_w, map_h)
    img.alpha_composite(basemap, (map_left, map_top))

    bbox_left_world, bbox_bottom_world = _latlon_to_world_pixels(min_lat, min_lon, zoom)
    bbox_right_world, bbox_top_world = _latlon_to_world_pixels(max_lat, max_lon, zoom)
    bbox_left = min(bbox_left_world, bbox_right_world)
    bbox_right = max(bbox_left_world, bbox_right_world)
    bbox_top = min(bbox_top_world, bbox_bottom_world)
    bbox_bottom = max(bbox_top_world, bbox_bottom_world)

    def to_panel_px(lat, lon):
        world_x, world_y = _latlon_to_world_pixels(lat, lon, zoom)
        px = ((world_x - bbox_left) / max(bbox_right - bbox_left, 1.0)) * map_w
        py = ((world_y - bbox_top) / max(bbox_bottom - bbox_top, 1.0)) * map_h
        return px, py

    def to_px(lat, lon):
        panel_x, panel_y = to_panel_px(lat, lon)
        return map_left + panel_x, map_top + panel_y

    overlay = Image.new("RGBA", (map_w, map_h), (0, 0, 0, 0))
    overlay_draw = ImageDraw.Draw(overlay)

    for ring in boundary_rings:
        if len(ring) < 3:
            continue
        points = [to_panel_px(lat, lon) for lon, lat in ring]
        closed_points = points + [points[0]]
        overlay_draw.polygon(points, fill=(59, 130, 246, 15), outline=None)

    # Rings are rendered on a separate alpha layer so road labels remain visible beneath them.
    placed_label_boxes = []

    def _rects_overlap(a, b, padding=0):
        return not (
            (a[2] + padding) <= b[0]
            or (b[2] + padding) <= a[0]
            or (a[3] + padding) <= b[1]
            or (b[3] + padding) <= a[1]
        )

    def _candidate_label_boxes(px, py, card_w, card_h):
        offsets = []
        for distance in (26, 52, 84):
            offsets.extend([
                (distance, -card_h - 14),
                (distance, 16),
                (-card_w - distance, -card_h - 14),
                (-card_w - distance, 16),
                (distance, -card_h - 52),
                (-card_w - distance, -card_h - 52),
            ])

        candidates = []
        for offset_x, offset_y in offsets:
            left = min(max(map_left + 12, px + offset_x), map_right - card_w - 12)
            top = min(max(map_top + 12, py + offset_y), map_bottom - card_h - 12)
            candidates.append((left, top, left + card_w, top + card_h))
        return candidates

    def _pick_label_box(px, py, card_w, card_h):
        candidates = _candidate_label_boxes(px, py, card_w, card_h)
        for rect in candidates:
            if all(not _rects_overlap(rect, other, padding=10) for other in placed_label_boxes):
                return rect
        return min(
            candidates,
            key=lambda rect: sum(
                max(0, min(rect[2], other[2]) - max(rect[0], other[0]))
                * max(0, min(rect[3], other[3]) - max(rect[1], other[1]))
                for other in placed_label_boxes
            ) if placed_label_boxes else 0,
        )

    for idx, site in enumerate(valid_sites, start=1):
        px, py = to_panel_px(site["latitude"], site["longitude"])
        east_edge = geodesic(meters=ring_radius_m).destination((site["latitude"], site["longitude"]), 90)
        ring_px, _ = to_panel_px(east_edge.latitude, east_edge.longitude)
        radius_px = abs(ring_px - px)
        overlay_draw.ellipse(
            [px - radius_px, py - radius_px, px + radius_px, py + radius_px],
            outline=(184, 55, 49, 210),
            width=5,
            fill=(217, 79, 67, 52),
        )

    img.alpha_composite(overlay, (map_left, map_top))

    if boundary_rings:
        boundary_draw = ImageDraw.Draw(img, "RGBA")
        for ring in boundary_rings:
            if len(ring) < 3:
                continue
            points = [to_panel_px(lat, lon) for lon, lat in ring]
            closed_points = points + [points[0]]
            _draw_dashed_line(boundary_draw, closed_points, color=(30, 64, 175, 255), width=3, dash_length=10, gap_length=6)

    for idx, site in enumerate(valid_sites, start=1):
        px, py = to_px(site["latitude"], site["longitude"])
        draw.ellipse([px - 12, py - 12, px + 12, py + 12], fill="#d92d20", outline="#ffffff", width=3)
        marker_text = str(idx)
        marker_box = draw.textbbox((0, 0), marker_text, font=marker_font)
        marker_w = marker_box[2] - marker_box[0]
        marker_h = marker_box[3] - marker_box[1]
        draw.text((px - marker_w / 2, py - marker_h / 2 - 1), marker_text, fill="#ffffff", font=marker_font)

        label = f"Site {idx}: {site_labels[idx - 1]}"
        text_bbox = draw.textbbox((0, 0), label, font=label_font)
        label_w = (text_bbox[2] - text_bbox[0]) + 24
        label_h = (text_bbox[3] - text_bbox[1]) + 18
        label_box = _pick_label_box(px, py, label_w, label_h)
        placed_label_boxes.append(label_box)
        label_x, label_y = label_box[0], label_box[1]
        anchor_x = max(label_x, min(px, label_x + label_w))
        anchor_y = max(label_y, min(py, label_y + label_h))
        draw.line([(px, py), (anchor_x, anchor_y)], fill="#7d9bb9", width=2)
        draw.rounded_rectangle(
            [label_x, label_y, label_x + label_w, label_y + label_h],
            radius=10,
            fill="#ffffff",
            outline="#b9cce0",
            width=2,
        )
        draw.text((label_x + 12, label_y + 9), label, fill="#16324f", font=label_font)

    if boundary_name and state:
        location_text = f"{boundary_name}, {state}"
    else:
        location_text = boundary_name or "Survey Area"
    draw.text((40, 30), "DFR SITE MAP", fill="#ffffff", font=title_font)
    draw.text(
        (40, 76),
        f"{location_text}   |   {len(valid_sites)} site(s)   |   2-mile operational rings",
        fill="#d8e4f2",
        font=subtitle_font,
    )

    legend_box = [width - 360, 24, width - 36, 96]
    draw.rounded_rectangle(legend_box, radius=14, fill="#173456", outline="#4c7096", width=2)
    draw.ellipse([legend_box[0] + 18, legend_box[1] + 16, legend_box[0] + 42, legend_box[1] + 40], fill="#d92d20", outline="#ffffff", width=2)
    draw.text((legend_box[0] + 52, legend_box[1] + 14), "Site marker", fill="#ffffff", font=small_font)
    draw.ellipse(
        [legend_box[0] + 18, legend_box[1] + 42, legend_box[0] + 42, legend_box[1] + 66],
        outline="#d94b45",
        width=3,
        fill="#f2bfba",
    )
    draw.text((legend_box[0] + 52, legend_box[1] + 40), "2-mile ring", fill="#ffffff", font=small_font)
    draw.rectangle([legend_box[0] + 170, legend_box[1] + 44, legend_box[0] + 194, legend_box[1] + 64], fill="#dcecff", outline="#2f6ea8", width=2)
    draw.text((legend_box[0] + 204, legend_box[1] + 40), "City boundary", fill="#ffffff", font=small_font)

    north_x = map_right - 42
    north_y = map_bottom - 42
    draw.polygon([(north_x, north_y - 30), (north_x - 12, north_y + 8), (north_x + 12, north_y + 8)], fill="#16324f")
    draw.text((north_x - 8, north_y + 12), "N", fill="#16324f", font=small_font)

    roster_width = 300
    roster_height = max(90, 42 + (len(valid_sites) * 24))
    roster_box = [map_right - roster_width - 28, map_top + 28, map_right - 28, map_top + 28 + roster_height]
    draw.rounded_rectangle(roster_box, radius=14, fill="#ffffff", outline="#b9cce0", width=2)
    draw.text((roster_box[0] + 18, roster_box[1] + 12), "Sites", fill="#16324f", font=label_font)
    for idx, site in enumerate(valid_sites, start=1):
        roster_y = roster_box[1] + 44 + ((idx - 1) * 24)
        roster_name = site_labels[idx - 1]
        draw.text((roster_box[0] + 18, roster_y), f"{idx}. {roster_name}", fill="#35516d", font=small_font)

    draw.text(
        (40, height - 48),
        "Street basemap: CARTO / OpenStreetMap. Rings are translucent for label readability. Boundary shown where available.",
        fill="#4c627a",
        font=small_font,
    )

    img.convert("RGB").save(output_map_path)
    return output_map_path

def create_engineering_drawing(bg_path, output_path, markers, engineer_note, address=None, overlay_image_path=None, image_placements=None, images_dir=None):
    """
    Renders an engineering drawing with a black tech sidebar legend (right),
    rounded callout labels with connecting path lines, and an engineer's note box (bottom left).
    Optionally includes placed images at specified coordinates.
    """
    # Canvas Size
    canvas_w, canvas_h = 1200, 675
    bg_w = 960  # Background image width
    sidebar_w = 240  # Sidebar width

    # Create master canvas
    canvas = Image.new('RGB', (canvas_w, canvas_h), color='#000000')
    draw = ImageDraw.Draw(canvas)
    
    # Initialize fonts (must be done before any text rendering below)
    font_bubble = _load_font(bold=True, size=28)
    font_legend = _load_font(bold=True, size=26)
    font_title = _load_font(bold=True, size=22)
    font_small = _load_font(bold=False, size=14)
    font_note = _load_font(bold=False, size=16)

    # Draw background image
    if os.path.exists(bg_path):
        try:
            with Image.open(bg_path) as img:
                img_resized = img.resize((bg_w, canvas_h))
                canvas.paste(img_resized, (0, 0))
        except Exception as e:
            draw.rectangle([0, 0, bg_w, canvas_h], fill='#1e293b')
            draw.text((100, 300), f"Error loading image: {e}", fill='#ffffff')
    else:
        draw.rectangle([0, 0, bg_w, canvas_h], fill='#1e293b')
        draw.text((100, 300), "No background image selected", fill='#ffffff')

    # Add Address at the top center of the photo area
    if address:
        # Draw address header at top center of photo area
        addr_text = _format_short_address(address)
        addr_bbox = draw.textbbox((0, 0), addr_text, font=font_title)
        addr_width = addr_bbox[2] - addr_bbox[0]
        addr_height = addr_bbox[3] - addr_bbox[1]

        # Position at top center of photo area, clamped to stay in bounds
        addr_x = max(12, (bg_w - addr_width) // 2)
        addr_y = 10

        # If text is wider than photo area, truncate with ellipsis
        if addr_width > bg_w - 30:
            while addr_width > bg_w - 30 and len(addr_text) > 10:
                addr_text = addr_text[:-4] + "..."
                addr_bbox = draw.textbbox((0, 0), addr_text, font=font_title)
                addr_width = addr_bbox[2] - addr_bbox[0]
                addr_height = addr_bbox[3] - addr_bbox[1]
            addr_x = max(12, (bg_w - addr_width) // 2)

        # Draw background box for readability
        padding = 12
        draw.rectangle(
            [addr_x - padding, addr_y - padding, addr_x + addr_width + padding, addr_y + addr_height + padding],
            fill='#000000',
            outline='#3B82F6',
            width=2
        )
        draw.text((addr_x, addr_y), addr_text, fill='#F8FAFC', font=font_title)

    # Add BRINC logo in the upper-left corner of the photo area.
    logo_path = os.path.join(os.path.dirname(__file__), "images", "BRINC_Logo_White.png")
    if os.path.exists(logo_path):
        try:
            with Image.open(logo_path) as logo_src:
                logo = logo_src.convert("RGBA")
                max_width = 180
                logo_ratio = max_width / float(logo.width)
                logo_size = (max_width, max(1, int(logo.height * logo_ratio)))
                logo = logo.resize(logo_size, Image.LANCZOS)

                # Remove the black background so the mark sits cleanly on the photo.
                pixels = logo.load()
                for y in range(logo.height):
                    for x in range(logo.width):
                        r, g, b, a = pixels[x, y]
                        if r < 20 and g < 20 and b < 20:
                            pixels[x, y] = (r, g, b, 0)
                        else:
                            pixels[x, y] = (r, g, b, 255)
                canvas.paste(logo, (20, 20), logo)
        except Exception as e:
            draw.text((20, 20), f"Logo load error: {e}", fill='#ffffff')

    # Draw Tech Sidebar Legend on the Right
    # Dark black background
    draw.rectangle([bg_w, 0, canvas_w, canvas_h], fill='#09090b')
    
    # Legend Items
    legend_items = [
        ("Ethernet", "#F97316"),   # Orange
        ("RF Xmit", "#A855F7"),    # Purple
        ("Electric", "#EF4444"),   # Red
        ("Unistrut", "#94A3B8"),   # Grey/Blue
        ("Lift", "#EAB308"),       # Yellow
        ("Station", "#F8FAFC")     # White
    ]
    
    # Draw "LEGEND" header
    legend_header_y = 25
    draw.text((bg_w + 20, legend_header_y), "LEGEND", fill='#F8FAFC', font=font_title)
    draw.line([(bg_w + 20, legend_header_y + 30), (canvas_w - 20, legend_header_y + 30)], fill='#334155', width=1)

    # Center legend items vertically in sidebar (below header, above brinc text)
    legend_item_spacing = 45
    legend_total_height = len(legend_items) * legend_item_spacing
    legend_start_y = legend_header_y + 50  # Start below header line

    for i, (title, color) in enumerate(legend_items):
        item_y = legend_start_y + i * legend_item_spacing
        # Draw colored circle swatch (20px diameter)
        cx, cy = bg_w + 30, item_y + 10
        draw.ellipse([cx - 10, cy - 10, cx + 10, cy + 10], fill=color)
        # Draw label text
        draw.text((bg_w + 55, item_y), title, fill=color, font=font_legend)

    # Draw "brinc" stylized logo at the bottom right
    draw.text((bg_w + 20, canvas_h - 60), "brinc", fill='#F8FAFC', font=font_legend)

    # Draw Markers & Callouts on the background image
    # Track placed bubble rectangles for collision avoidance
    placed_bubbles = []

    for m in markers:
        # Convert percent to actual pixel coords
        nx = int((m['node_x'] / 100.0) * bg_w)
        ny = int((m['node_y'] / 100.0) * canvas_h)

        m_type = m.get('type', 'Electric')
        label_text = m.get('label', '')

        # Select color matching type
        color_map = {
            'Data': '#F97316',
            'RF': '#A855F7',
            'Unistrut': '#94A3B8',
            'Lift': '#EAB308',
        }
        border_color = color_map.get(m_type, '#EF4444')

        # Measure actual text size for accurate bubble dimensions
        bbox = draw.textbbox((0, 0), label_text, font=font_bubble)
        text_w = bbox[2] - bbox[0]
        text_h = bbox[3] - bbox[1]

        pad_x, pad_y = 14, 14
        bubble_w = text_w + pad_x * 2
        bubble_h = text_h + pad_y * 2

        # Smart bubble placement: try 4 directions, check edge + collision
        offset_x = int(bg_w * 0.15)
        offset_y = int(canvas_h * 0.08)

        candidates = [
            (nx - offset_x - bubble_w // 2, ny - offset_y - bubble_h // 2),  # upper-left
            (nx + offset_x - bubble_w // 2, ny - offset_y - bubble_h // 2),  # upper-right
            (nx + offset_x - bubble_w // 2, ny + offset_y - bubble_h // 2),  # lower-right
            (nx - offset_x - bubble_w // 2, ny + offset_y - bubble_h // 2),  # lower-left
        ]

        best_pos = None
        best_overlap = float('inf')

        for cx, cy in candidates:
            # Clamp to photo area bounds
            cx = max(2, min(cx, bg_w - bubble_w - 2))
            cy = max(2, min(cy, canvas_h - bubble_h - 2))
            candidate_rect = (cx, cy, cx + bubble_w, cy + bubble_h)

            # Check if bubble fits within photo area
            fits_edge = (cx >= 0 and cy >= 0 and cx + bubble_w <= bg_w and cy + bubble_h <= canvas_h)

            # Calculate total overlap with already-placed bubbles
            total_overlap = 0
            for pb in placed_bubbles:
                # AABB intersection area
                ox1 = max(candidate_rect[0], pb[0])
                oy1 = max(candidate_rect[1], pb[1])
                ox2 = min(candidate_rect[2], pb[2])
                oy2 = min(candidate_rect[3], pb[3])
                if ox1 < ox2 and oy1 < oy2:
                    total_overlap += (ox2 - ox1) * (oy2 - oy1)

            if fits_edge and total_overlap == 0:
                best_pos = (cx, cy)
                break
            elif total_overlap < best_overlap:
                best_overlap = total_overlap
                best_pos = (cx, cy)

        if best_pos is None:
            best_pos = (candidates[0][0], candidates[0][1])

        lx, ly = best_pos[0] + bubble_w // 2, best_pos[1] + bubble_h // 2
        rx1, ry1 = best_pos
        rx2, ry2 = rx1 + bubble_w, ry1 + bubble_h

        # Track this bubble for future collision checks
        placed_bubbles.append((rx1, ry1, rx2, ry2))

        # Draw connection line from node dot to bubble center
        draw.line([(nx, ny), (lx, ly)], fill=border_color, width=3)

        # Draw node dot and bubble endpoint dot
        draw.ellipse([nx - 6, ny - 6, nx + 6, ny + 6], fill=border_color)
        draw.ellipse([lx - 4, ly - 4, lx + 4, ly + 4], fill=border_color)

        # Draw rounded rectangle bubble
        draw.rounded_rectangle([rx1, ry1, rx2, ry2], radius=12, outline=border_color, width=3, fill='#000000')

        # Draw centered text inside bubble
        text_x = rx1 + (bubble_w - text_w) // 2
        text_y = ry1 + (bubble_h - text_h) // 2
        draw.text((text_x, text_y), label_text, fill='#FFFFFF', font=font_bubble)

    # Draw Engineer's Note Box centered in the sidebar
    note_w, note_h = 220, 250
    nx1 = bg_w + (sidebar_w - note_w) // 2
    ny1 = canvas_h - note_h - 15
    nx2, ny2 = nx1 + note_w, ny1 + note_h

    draw.rectangle([nx1, ny1, nx2, ny2], fill='#FFFFFF', outline='#000000', width=4)
    draw.text((nx1 + 10, ny1 + 10), "ENGINEER'S NOTE:\n-----------------------------", fill='#000000', font=font_note)
    
    # Wrap note text manually
    words = engineer_note.split()
    lines = []
    current_line = []
    for word in words:
        if len(" ".join(current_line + [word])) * 8 > (note_w - 20):
            lines.append(" ".join(current_line))
            current_line = [word]
        else:
            current_line.append(word)
    if current_line:
        lines.append(" ".join(current_line))
        
    line_y = ny1 + 40
    for line in lines[:10]: # Cap at 10 lines
        draw.text((nx1 + 10, line_y), line, fill='#000000', font=font_small)
        line_y += 18

    # Draw placed images at their coordinates
    if image_placements and images_dir:
        for placement in image_placements:
            try:
                image_name = placement.get('image_name')
                x_pct = placement.get('x', 0)
                y_pct = placement.get('y', 0)

                image_path = os.path.join(images_dir, image_name)
                if os.path.exists(image_path):
                    with Image.open(image_path) as img_src:
                        img_overlay = img_src.convert("RGBA")
                        # Scale image to reasonable size (max 100px on largest dimension)
                        max_size = 100
                        img_overlay.thumbnail((max_size, max_size), Image.LANCZOS)

                        # Convert percentages to actual pixel coordinates
                        img_x = int((x_pct / 100.0) * bg_w) - (img_overlay.width // 2)
                        img_y = int((y_pct / 100.0) * canvas_h) - (img_overlay.height // 2)

                        # Ensure image stays within bounds
                        img_x = max(0, min(img_x, bg_w - img_overlay.width))
                        img_y = max(0, min(img_y, canvas_h - img_overlay.height))

                        canvas.paste(img_overlay, (img_x, img_y), img_overlay)
            except Exception as e:
                print(f"Error placing image {placement.get('image_name')}: {e}")

    # Add overlay image in bottom-right corner if provided
    if overlay_image_path and os.path.exists(overlay_image_path):
        try:
            with Image.open(overlay_image_path) as overlay_src:
                overlay = overlay_src.convert("RGBA")
                # Scale overlay to reasonable size (max 120px on largest dimension)
                max_size = 120
                overlay.thumbnail((max_size, max_size), Image.LANCZOS)

                # Position in bottom-right corner with padding
                padding = 15
                overlay_x = canvas_w - overlay.width - padding
                overlay_y = canvas_h - overlay.height - padding

                # Paste with alpha transparency
                canvas.paste(overlay, (overlay_x, overlay_y), overlay)
        except Exception as e:
            print(f"Error adding overlay image: {e}")

    canvas.save(output_path)
    return output_path

def add_styled_table(doc, data, headers):
    """
    Construct a clean, polished table with a strong header row and a highlighted label column.
    Supports 2 or 3 columns depending on len(headers).
    """
    ncols = len(headers)
    table = doc.add_table(rows=len(data) + 1, cols=ncols)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    for r_idx, row in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_val in enumerate(row):
            row_cells[c_idx].text = str(cell_val)

    if ncols == 2:
        widths = [Inches(2.25), Inches(4.25)]
    elif ncols == 3:
        widths = [Inches(2.20), Inches(3.00), Inches(1.30)]
    else:
        widths = [Inches(6.5 / ncols)] * ncols
    _style_polished_table(table, widths=widths)
    doc.add_paragraph()


def _get_contact_value(customer_info, key, default="DNA"):
    value = ""
    if customer_info:
        value = customer_info.get(key, "")
    return value if value not in ("", None) else default


def _format_contact_block(*parts):
    lines = [str(part).strip() for part in parts if part not in ("", None)]
    return "\n".join(lines) if lines else "DNA"


def _find_contact_row(customer_info, role_name):
    """Return the first structured contact row matching a role name."""
    for row in (customer_info or {}).get("contacts", []):
        if str(row.get("role", "")).strip().lower() == role_name.strip().lower():
            return row
    return {}


def _format_role_contact(customer_info, role_name, name_key, email_key, phone_key=None, default="DNA"):
    """Prefer structured contact rows, then flat legacy keys, then DNA."""
    customer_info = customer_info or {}
    contact = _find_contact_row(customer_info, role_name)
    name = (contact.get("name") or customer_info.get(name_key) or "").strip()
    email = (contact.get("email") or customer_info.get(email_key) or "").strip()
    phone = ""
    if phone_key:
        phone = (contact.get("phone") or customer_info.get(phone_key) or "").strip()

    lines = [value for value in (name, email, phone) if value]
    return "\n".join(lines) if lines else default


def _format_scalar_contact(customer_info, key, default="DNA"):
    """Return a single-line scalar contact value from the structured or legacy payload."""
    customer_info = customer_info or {}
    value = str(customer_info.get(key, "") or "").strip()
    if not value:
        for row in customer_info.get("contacts", []):
            if str(row.get("role", "")).strip().lower() == key.replace("_", " ").lower():
                value = str(row.get("name", "") or "").strip()
                break
    return value if value else default


def _normalized_contact_rows(customer_info):
    """Return contact rows for reporting, defaulting missing roles to Other."""
    rows = []
    for row in (customer_info or {}).get("contacts", []):
        role = str(row.get("role", "") or "").strip() or "Other"
        name = str(row.get("name", "") or "").strip()
        title = str(row.get("title", "") or "").strip()
        email = str(row.get("email", "") or "").strip()
        phone = str(row.get("phone", "") or "").strip()
        if not any((role, name, title, email, phone)):
            continue
        rows.append((role, name or "DNA", title or "DNA", email or "DNA", phone or "DNA"))
    return rows


def _add_contacts_table(doc, customer_info):
    """Append the full contact roster so non-primary roles are preserved in the report."""
    rows = _normalized_contact_rows(customer_info)
    if not rows:
        return

    doc.add_paragraph().add_run("POINTS OF CONTACT").bold = True
    table = doc.add_table(rows=len(rows) + 1, cols=5)
    table.style = "Table Grid"

    headers = ["Role", "Name", "Title / Rank", "Email", "Phone"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = value

    _style_polished_table(
        table,
        widths=[Inches(1.0), Inches(1.45), Inches(1.6), Inches(1.75), Inches(1.1)],
        header_fill="FFFFFF",
    )
    for row in table.rows:
        for cell in row.cells:
            _style_table_cell(cell, bold=False, font_size=10, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()


def _add_poc_table(doc, customer_info):
    """Build a polished 2-column general information table that mirrors the reference document."""
    customer_info = customer_info or {}

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_run = h.add_run("GENERAL")
    h_run.bold = True
    h_run.font.name = "Aptos"
    h_run.font.size = Pt(22)
    h_run.font.color.rgb = RGBColor(0, 0, 0)

    rows = [
        ("Agency Name\nAgency Address", _format_contact_block(
            _get_contact_value(customer_info, "agency_name", ""),
            _get_contact_value(customer_info, "agency_address", ""),
        )),
        ("Point of Contact\n(Name, E-Mail, Phone #)", _format_contact_block(
            _get_contact_value(customer_info, "poc_name", ""),
            _get_contact_value(customer_info, "poc_email", ""),
            _get_contact_value(customer_info, "poc_phone", ""),
        )),
        ("RTCC/RTIC\n(Name, E-Mail, Phone #)", _format_contact_block(
            *_format_role_contact(customer_info, "RTCC", "rtcc_name", "rtcc_email", "rtcc_phone").split("\n")
        )),
        ("Information Technology\n(Name, E-Mail, Phone #)", _format_contact_block(
            *_format_role_contact(customer_info, "IT", "it_director", "it_email", "it_phone").split("\n")
        )),
        ("Facilities Engineer\n(Name, E-Mail, Phone #)", _format_contact_block(
            *_format_role_contact(customer_info, "Facilities", "facilities_engineer", "facilities_email", "facilities_phone").split("\n")
        )),
        ("Radio Shop Engineer\n(Name, E-Mail, Phone #)", _format_contact_block(
            *_format_role_contact(customer_info, "Radio Shop", "radio_shop_name", "radio_shop_email", "radio_shop_phone").split("\n")
        )),
        ("Crane Contractor", _format_scalar_contact(customer_info, "crane_contractor", "DNA")),
        ("Tower Climber Contractor", _format_scalar_contact(customer_info, "tower_climber_contractor", "DNA")),
        ("BRINC Project Manager", _format_scalar_contact(customer_info, "brinc_pm", "DNA")),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    for row_idx, (label, value) in enumerate(rows):
        table.rows[row_idx].cells[0].text = label
        table.rows[row_idx].cells[1].text = value

    _style_polished_table(table, widths=[Inches(2.55), Inches(4.0)], header_fill="FFFFFF")
    for row in table.rows:
        _set_cell_shading(row.cells[0], "C7D9F1")
        _set_cell_shading(row.cells[1], "FFFFFF")
        _style_table_cell(row.cells[0], bold=False, font_size=11, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.LEFT)
        _style_table_cell(row.cells[1], bold=False, font_size=11, color=RGBColor(0, 0, 0), align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()
    _add_contacts_table(doc, customer_info)


def generate_word_report(site_data_list, output_filepath, customer_info=None, drive_manager=None, drive_reports_folder_id=None, progress_callback=None):
    """
    Generate the Site Survey Word Document.
    site_data_list items contain pre-populated 'agency_name' field from processor.

    If drive_manager and drive_reports_folder_id are provided, the report
    will also be uploaded to Google Drive.

    Args:
        progress_callback: Optional callable(step_name: str) called before each major step.
    """
    def _progress(step):
        if progress_callback:
            progress_callback(step)

    _progress("Initializing report document...")
    doc = Document()
    _add_report_footer(doc)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Aptos'
    font.size = Pt(10)

    if not customer_info:
        customer_info = {
            "agency_name": "",
            "agency_address": "",
            "poc_name": "",
            "poc_email": "",
            "poc_phone": "",
            "it_director": "",
            "it_email": "",
            "facilities_engineer": "",
            "facilities_email": "",
            "contacts": [],
            "report_date": datetime.date.today().strftime("%B %d, %Y"),
        }
    survey_date = _extract_survey_date(site_data_list=site_data_list, customer_info=customer_info)
    if survey_date:
        customer_info["survey_date"] = survey_date
        customer_info["report_date"] = _format_report_date(survey_date)
    report_date = customer_info.get("report_date", _format_report_date(survey_date))
    surveyor = customer_info.get("surveyor", "")
    display_agency = (
        customer_info.get("jurisdiction_name")
        or customer_info.get("agency_name")
        or ""
    )
        
    p_title = doc.add_paragraph()
    p_title.alignment = 1
    title_run = p_title.add_run("DFR SITE SURVEY REPORT")
    title_run.bold = True
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(29, 59, 103)
    
    p_agency = doc.add_paragraph()
    p_agency.alignment = 1
    agency_run = p_agency.add_run(
        f"Agency: {display_agency}\n"
        f"Survey Date: {report_date}\n"
        f"Surveyor: {surveyor}"
    )
    agency_run.font.name = "Aptos"
    agency_run.font.size = Pt(15)
    agency_run.font.color.rgb = RGBColor(98, 98, 98)
    
    # 1. Map Visualisation Section
    _progress("Generating site map visualization...")
    doc.add_paragraph().add_run("Site Detail & Map Visualisation").bold = True
    map_image_path = os.path.join(os.path.dirname(output_filepath), "dfr_site_map.png")
    draw_styled_map(site_data_list, map_image_path)
    if os.path.exists(map_image_path):
        doc.add_picture(map_image_path, width=Inches(6.0))
        doc.add_paragraph()
        
    # 2. Installation Timeframe
    delivery_target = customer_info.get("survey_delivery_target", "TBD")
    follow_up = customer_info.get("follow_up_requirements", "") or "Infrastructure checklist completion prior to hardware delivery"
    action_items = customer_info.get("action_items", "") or "Confirm ethernet and dedicated power connectivity is established 30 days before installation."
    timeframe_data = [
        ("Survey / Delivery Target", delivery_target),
        ("Follow up requirements", follow_up),
        ("Action items", action_items),
    ]
    add_styled_table(doc, timeframe_data, ["INSTALLATION TIMEFRAME", "NOTES / VALUE"])
    
    doc.add_page_break()
    
    # 4. Loop through individual site nodes
    for idx, site in enumerate(site_data_list):
        _progress(f"Building site {idx+1}/{len(site_data_list)}: {site['address'].split(',')[0]}...")
        site_name = site['address'].split(',')[0]
        analysis = site.get('analysis', {})
        # Note: Each site includes pre-populated 'agency_name' and 'city' fields from processor
        site_agency = site.get('agency_name', 'Police Department')
        
        h_site = doc.add_paragraph()
        h_site_run = h_site.add_run(f"Location #{idx+1}/{len(site_data_list)} - {site_name.upper()}")
        h_site_run.bold = True
        h_site_run.font.size = Pt(14)
        h_site_run.font.color.rgb = RGBColor(0, 48, 96)
        
        # If engineering layout has been generated for this site, embed it
        eng_drawing_path = os.path.join(site.get('folder_path', '.'), "engineering_layout.png")
        if os.path.exists(eng_drawing_path):
            doc.add_paragraph().add_run("Engineering Layout Drawing").bold = True
            try:
                doc.add_picture(eng_drawing_path, width=Inches(6.0))
            except Exception as e:
                doc.add_paragraph(f"[Error embedding engineering drawing: {e}]")
            doc.add_paragraph()

        # Site Details
        building_height = site.get('building_height')
        if building_height:
            height_str = f"{building_height:.0f} ft" if isinstance(building_height, (int, float)) else str(building_height)
        else:
            height_str = "Assessment required"

        elevation = site.get('elevation')
        elevation_str = ""
        if elevation:
            if isinstance(elevation, (int, float)):
                elevation_str = f"{elevation:.0f} ft"
            else:
                elevation_str = str(elevation)

        details_data = [
            ("Site Name", site_name),
            ("Site Address", _format_short_address(site['address'])),
            ("Height of building", height_str),
        ]
        if elevation_str:
            details_data.append(("Ground Elevation", elevation_str))
        details_data.extend([
            ("Access to roof", analysis.get('roof_access', 'Unknown')),
            ("Roof type", analysis.get('roof_type', 'Unknown'))
        ])
        add_styled_table(doc, details_data, ["SITE DETAILS", "NOTES / VALUE"])
        
        # Considerations & Airspace
        op_notes = site.get('operational_considerations', "Coordinate install with local facilities team. Clear line of sight required.")
        considerations_data = [
            ("Operational Considerations", op_notes),
            ("Airspace Class", site.get('airspace', 'Unknown')),
            ("Distance to nearest airfield", site.get('airfield_info', 'Unknown'))
        ]
        add_styled_table(doc, considerations_data, ["OPERATIONAL & AIRSPACE", "NOTES / VALUE"])

        # Deployment Requirements
        power_req = customer_info.get("power_circuit_requirements", "120V / 15A Dedicated Circuit, Outdoor Rated") if customer_info else "120V / 15A Dedicated Circuit, Outdoor Rated"
        network_req = customer_info.get("internet_ethernet_access", "DHCP on isolated VLAN (unrestricted outbound)") if customer_info else "DHCP on isolated VLAN (unrestricted outbound)"
        deployment_data = [
            ("Location (Lat/Long)", f"{site['latitude']:.6f}, {site['longitude']:.6f}"),
            ("Mount Placement Type", "Rooftop / Parapet Mount preferred" if "hatch" in str(analysis).lower() else "Ground Sled / Ballasted"),
            ("Raised Platform required (Snow)", "Yes" if "TPO" in str(analysis) or "EPDM" in str(analysis) else "No"),
            ("Emergency Landing Zone", "Yes (Debris-free zone verified)"),
            ("Power Circuit Requirements", power_req),
            ("Internet / Ethernet Access", network_req)
        ]
        add_styled_table(doc, deployment_data, ["DEPLOYMENT SPECIFICATIONS", "NOTES / VALUE"])
        
        # Add ONLY selected images
        selected_images = [img for img in site.get('images', []) if img.get('selected_for_report', True)]
        if selected_images:
            _progress(f"Embedding {len(selected_images)} photo(s) for site {idx+1}...")
            p_photos = doc.add_paragraph()
            p_photos.add_run("Survey Photos & Visual Evidence").bold = True
            for img in selected_images:
                img_path = img.get('dest_path') or img.get('path')
                if os.path.exists(img_path):
                    doc.add_paragraph(f"Photo: {img.get('filename')}")
                    try:
                        doc.add_picture(img_path, width=Inches(3.5))
                    except Exception as e:
                        doc.add_paragraph(f"[Unable to embed photo: {e}]")
                    doc.add_paragraph()
                    
        doc.add_page_break()
        
    _progress("Saving report document...")
    doc.save(output_filepath)

    # If using Google Drive, upload report
    if drive_manager and drive_reports_folder_id:
        _progress("Uploading report to Google Drive...")
        try:
            report_filename = os.path.basename(output_filepath)
            drive_manager.upload_file(
                output_filepath,
                drive_reports_folder_id,
                file_name=report_filename
            )
        except Exception as e:
            print(f"Failed to upload report to Google Drive: {e}")

    return output_filepath


def _yn(value):
    """Convert bool/None to Yes/No/— string."""
    if value is True:
        return "Yes"
    elif value is False:
        return "No"
    return "—"


def generate_candidate_site_report(candidate_sites, output_filepath,
                                    customer_info=None, drive_manager=None,
                                    drive_reports_folder_id=None,
                                    progress_callback=None):
    """Generate a dynamic DOCX report from CandidateSite objects.

    Report structure:
    1. Executive Summary
    2. Candidate Site sections (1-N, dynamic)
    3. Installer Quick Reference (1 page per site)
    4. Annotated Photo Appendix

    Args:
        progress_callback: Optional callable(step_name: str) called before each major step.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def _progress(step):
        if progress_callback:
            progress_callback(step)

    _progress("Initializing report document...")
    doc = Document()
    _add_report_footer(doc)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # ── 1. Executive Summary ──
    title = doc.add_heading("DFR SITE SURVEY REPORT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    agency = ""
    if candidate_sites:
        agency = candidate_sites[0].identity.agency_name
    survey_date = _extract_survey_date(candidate_sites=candidate_sites, customer_info=customer_info)
    surveyor = candidate_sites[0].identity.surveyor if candidate_sites else ""
    if customer_info:
        agency = customer_info.get("agency_name", agency)
        if survey_date:
            customer_info["survey_date"] = survey_date
            customer_info["report_date"] = _format_report_date(survey_date)
        surveyor = customer_info.get("surveyor", surveyor)

    _progress("Generating site map visualization...")
    doc.add_paragraph().add_run("Site Detail & Map Visualisation").bold = True
    map_image_path = os.path.join(os.path.dirname(output_filepath), "dfr_site_map.png")
    draw_styled_map(_survey_site_records(candidate_sites=candidate_sites), map_image_path)
    if os.path.exists(map_image_path):
        doc.add_picture(map_image_path, width=Inches(6.0))
        doc.add_paragraph()

    doc.add_paragraph(f"Agency: {agency}")
    doc.add_paragraph(f"Survey Date: {_format_report_date(survey_date)}")
    doc.add_paragraph(f"Surveyor: {surveyor}")
    doc.add_paragraph(f"Candidate Sites Found: {len(candidate_sites)}")

    doc.add_page_break()

    # ── 2. Candidate Site Sections ──
    for idx, site in enumerate(candidate_sites, start=1):
        _progress(f"Building site {idx}/{len(candidate_sites)}: {site.identity.site_name}...")
        doc.add_heading(f"Candidate Site {idx}: {site.identity.site_name}", level=1)

        # 2a. Site Overview
        doc.add_heading("Site Overview", level=2)
        overview_data = [
            ["Site ID", site.identity.site_id],
            ["Address", site.identity.site_address],
            ["Latitude", str(site.identity.site_latitude)],
            ["Longitude", str(site.identity.site_longitude)],
            ["Elevation", f"{site.identity.site_elevation} ft" if site.identity.site_elevation else "—"],
            ["Building Height", f"{site.structure.building_height} ft" if site.structure.building_height else "—"],
            ["Roof Type", site.structure.roof_type or "—"],
        ]
        add_styled_table(doc, overview_data, ["Field", "Value"])

        # 2b. Site Photos by Category
        _progress(f"Embedding photos for site {idx}...")
        doc.add_heading("Site Photos", level=2)
        for cat in ["Site", "Installation", "Infrastructure", "RF", "Access"]:
            cat_photos = [p for p in site.photos if p.category == cat and p.selected_for_report]
            if not cat_photos:
                continue
            doc.add_heading(f"{cat} Photos", level=3)
            for photo in cat_photos:
                if os.path.exists(photo.file_path):
                    try:
                        doc.add_picture(photo.file_path, width=Inches(5.0))
                    except Exception:
                        doc.add_paragraph(f"[Photo: {photo.photo_id}]")
                else:
                    doc.add_paragraph(f"[Photo: {photo.photo_id}]")

        # 2c. Checklist Summary
        doc.add_heading("Checklist Summary", level=2)
        summary_rows = []
        field_labels = {
            "ACCESS_TYPE": "Access Type", "ROOF_ACCESS": "Roof Access",
            "ESCORT_REQUIRED": "Escort Required", "KEY_REQUIRED": "Key Required",
            "BUILDING_HEIGHT": "Building Height", "ROOF_TYPE": "Roof Type",
            "PARAPET_HEIGHT": "Parapet Height", "ROOF_CONDITION": "Roof Condition",
            "POWER_AVAILABLE": "Power Available", "VOLTAGE_AVAILABLE": "Voltage",
            "DEDICATED_CIRCUIT": "Dedicated Circuit", "PANEL_LOCATION": "Panel Location",
            "ISP_PROVIDER": "ISP Provider", "DOWNLOAD_SPEED": "Download Speed",
            "UPLOAD_SPEED": "Upload Speed", "STATIC_IP_AVAILABLE": "Static IP",
            "LINE_OF_SIGHT_STATUS": "Line of Sight", "COVERAGE_DIRECTION": "Coverage Direction",
            "AIRSPACE_CLASS": "Airspace Class", "SITE_ELEVATION": "Elevation",
            "COUNTY_NAME": "County", "STATE_NAME": "State", "ZIP_CODE": "ZIP Code",
        }
        all_fields = site.to_csv_row()
        for field_id, label in field_labels.items():
            value = all_fields.get(field_id)
            if value is not None and value != "" and str(value) != "None":
                source = site.checklist_provenance.get(field_id, "pm")
                summary_rows.append([label, str(value), "Auto" if source == "auto" else "Manual"])
        if summary_rows:
            add_styled_table(doc, summary_rows, ["Field", "Value", "Source"])
        else:
            doc.add_paragraph("No checklist fields have been filled.")

        # 2d. Access
        doc.add_heading("Access", level=3)
        add_styled_table(doc, [
            ["Access Type", site.access.access_type or "—"],
            ["Roof Access", site.access.roof_access or "—"],
            ["Escort Required", _yn(site.access.escort_required)],
            ["Key Required", _yn(site.access.key_required)],
            ["After Hours Access", _yn(site.access.after_hours_access)],
            ["Parking Available", _yn(site.access.parking_available)],
        ], ["Field", "Value"])

        # 2e. Electrical Assessment
        doc.add_heading("Electrical Assessment", level=3)
        add_styled_table(doc, [
            ["Power Available", _yn(site.electrical.power_available)],
            ["Voltage", site.electrical.voltage_available or "—"],
            ["Breaker Available", _yn(site.electrical.breaker_available)],
            ["Dedicated Circuit", _yn(site.electrical.dedicated_circuit)],
            ["Panel Location", site.electrical.panel_location or "—"],
            ["Distance to Power", f"{site.electrical.distance_to_power} ft" if site.electrical.distance_to_power else "—"],
        ], ["Field", "Value"])

        # 2f. Network Assessment
        doc.add_heading("Network Assessment", level=3)
        add_styled_table(doc, [
            ["ISP Provider", site.network.isp_provider or "—"],
            ["Connection Type", site.network.connection_type or "—"],
            ["Download Speed", site.network.download_speed or "—"],
            ["Upload Speed", site.network.upload_speed or "—"],
            ["Static IP", _yn(site.network.static_ip_available)],
            ["Switch Location", site.network.switch_location or "—"],
            ["Distance to Network", f"{site.network.distance_to_network} ft" if site.network.distance_to_network else "—"],
        ], ["Field", "Value"])

        # 2g. RF Assessment
        doc.add_heading("RF Assessment", level=3)
        add_styled_table(doc, [
            ["Line of Sight", site.rf.line_of_sight_status or "—"],
            ["Obstructions - Trees", _yn(site.rf.obstruction_trees)],
            ["Obstructions - Buildings", _yn(site.rf.obstruction_buildings)],
            ["Obstructions - Water Towers", _yn(site.rf.obstruction_water_towers)],
            ["Obstructions - Cell Towers", _yn(site.rf.obstruction_cell_towers)],
            ["Coverage Direction", site.rf.coverage_direction or "—"],
        ], ["Field", "Value"])

        # 2h. Airspace Assessment
        doc.add_heading("Airspace Assessment", level=3)
        add_styled_table(doc, [
            ["Airspace Class", site.flight.airspace_class or "—"],
            ["Nearby Airports", site.flight.nearby_airports or "—"],
            ["Nearby Heliports", site.flight.nearby_heliports or "—"],
            ["Flight Restrictions", site.flight.flight_restrictions or "—"],
        ], ["Field", "Value"])

        # 2i. Data Quality / Provenance
        failed_lookups = [
            (field_id, label) for field_id, label in {
                "COUNTY_NAME": "Geocode (county/state/zip)",
                "SITE_ELEVATION": "Elevation lookup",
                "BUILDING_HEIGHT": "Building height estimate",
                "NEARBY_AIRPORTS": "Airport distance lookup",
                "NEARBY_HELIPORTS": "Heliport distance lookup",
            }.items()
            if site.checklist_provenance.get(field_id) == "failed"
        ]
        if failed_lookups:
            doc.add_heading("Data Quality Notes", level=3)
            for _, label in failed_lookups:
                doc.add_paragraph(f"  {label} — lookup failed, value unavailable", style="List Bullet")

        doc.add_page_break()

    # ── 3. Installer Quick Reference ──
    _progress("Building installer quick reference...")
    doc.add_heading("Installer Quick Reference", level=1)
    for idx, site in enumerate(candidate_sites, start=1):
        doc.add_heading(f"Site {idx}: {site.identity.site_name}", level=2)
        add_styled_table(doc, [
            ["Address", site.identity.site_address],
            ["Access Method", site.access.roof_access or site.access.access_type or "—"],
            ["Power Location", site.electrical.panel_location or "—"],
            ["Network Location", site.network.switch_location or "—"],
            ["Antenna Location", f"{site.rf.antenna_latitude}, {site.rf.antenna_longitude}" if site.rf.antenna_latitude else "—"],
            ["Escort Required", _yn(site.access.escort_required)],
        ], ["Item", "Details"])
        if customer_info:
            poc = customer_info.get("poc_name", "")
            phone = customer_info.get("poc_phone", "")
            if poc:
                doc.add_paragraph(f"Site Contact: {poc}  {phone}")
        doc.add_page_break()

    # ── 4. Annotated Photo Appendix ──
    _progress("Adding annotated photo appendix...")
    doc.add_heading("Annotated Photo Appendix", level=1)
    has_annotations = False
    for idx, site in enumerate(candidate_sites, start=1):
        folder = site.folder_path
        if folder and os.path.isdir(folder):
            for fname in sorted(os.listdir(folder)):
                if fname.startswith("engineering_layout") and fname.endswith(".png"):
                    img_path = os.path.join(folder, fname)
                    doc.add_heading(f"Site {idx}: {site.identity.site_name}", level=3)
                    try:
                        doc.add_picture(img_path, width=Inches(6.0))
                        has_annotations = True
                    except Exception:
                        doc.add_paragraph(f"[Annotation: {fname}]")
    if not has_annotations:
        doc.add_paragraph("No annotated engineering layouts available.")

    _progress("Saving report document...")
    doc.save(output_filepath)

    if drive_manager and drive_reports_folder_id:
        _progress("Uploading report to Google Drive...")
        try:
            drive_manager.upload_file(output_filepath, drive_reports_folder_id)
        except Exception:
            pass

    return output_filepath
